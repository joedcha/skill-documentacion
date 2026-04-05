#!/usr/bin/env python3
"""
md_to_docx.py — Convertidor Markdown → Word (.docx) con diseño profesional
Skill: /documentacion v2.0.0
"""

import sys
import os
import re
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx no está instalado.")
    print("Instalar con: pip install python-docx")
    sys.exit(1)


# ─── Utilidades de color ────────────────────────────────────────────────────

def hex_to_rgb(hex_color: str) -> tuple:
    """Convierte color hex a tupla RGB."""
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return (0, 229, 255)
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color: str):
    """Establece color de fondo de celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Establece bordes de celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), kwargs.get("val", "single"))
        border.set(qn("w:sz"), kwargs.get("sz", "4"))
        border.set(qn("w:color"), kwargs.get("color", "E0E0E0"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


# ─── Creación de portada ─────────────────────────────────────────────────────

def add_cover_page(doc, title: str, subtitle: str, company: str,
                   accent_hex: str, logo_path: str = None):
    """Agrega página de portada profesional."""
    accent_rgb = hex_to_rgb(accent_hex)
    dark_navy = "1A1A2E"

    # Logo (si se proporciona y existe)
    if logo_path and os.path.exists(logo_path):
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = logo_para.runs[0] if logo_para.runs else logo_para.add_run()
            run.add_picture(logo_path, width=Inches(1.8))
            logo_para.space_after = Pt(40)
        except Exception:
            pass

    # Espaciado superior
    for _ in range(4):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Línea de acento superior
    accent_para = doc.add_paragraph()
    accent_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    accent_run = accent_para.add_run("█" * 6)
    accent_run.font.color.rgb = RGBColor(*accent_rgb)
    accent_run.font.size = Pt(14)
    accent_para.space_after = Pt(12)

    # Título principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(title.upper())
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(*hex_to_rgb(dark_navy))
    title_para.space_after = Pt(8)

    # Subtítulo
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.name = "Calibri"
        sub_run.font.size = Pt(15)
        sub_run.font.color.rgb = RGBColor(80, 80, 100)
        sub_para.space_after = Pt(48)

    # Línea divisoria de acento
    div_para = doc.add_paragraph()
    div_run = div_para.add_run("─" * 52)
    div_run.font.color.rgb = RGBColor(*accent_rgb)
    div_run.font.size = Pt(10)
    div_para.space_after = Pt(16)

    # Empresa / desarrollador
    if company:
        comp_para = doc.add_paragraph()
        comp_run = comp_para.add_run(company)
        comp_run.font.name = "Calibri"
        comp_run.font.size = Pt(12)
        comp_run.font.color.rgb = RGBColor(100, 100, 120)
        comp_para.space_after = Pt(4)

    # Fecha
    MONTHS_ES = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    now = datetime.now()
    date_str = f"{now.day} de {MONTHS_ES[now.month]} de {now.year}"
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date_str)
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(140, 140, 160)

    doc.add_page_break()


# ─── Parser de Markdown ──────────────────────────────────────────────────────

def apply_inline_formatting(run_text: str, paragraph, font_name: str,
                             font_size: float, accent_rgb: tuple,
                             base_color: tuple = (30, 30, 50)):
    """Aplica formato inline: **negrita**, `código`, _cursiva_."""
    # Patrón para capturar formatos inline
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`|_[^_]+_|\*[^*]+\*)'
    parts = re.split(pattern, run_text)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)

        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
            run.font.color.rgb = RGBColor(*base_color)
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(font_size - 1)
            run.font.color.rgb = RGBColor(*accent_rgb)
        elif (part.startswith("_") and part.endswith("_")) or \
             (part.startswith("*") and part.endswith("*")):
            run.text = part[1:-1]
            run.italic = True
            run.font.color.rgb = RGBColor(*base_color)
        else:
            run.text = part
            run.font.color.rgb = RGBColor(*base_color)


def parse_table(doc, table_lines: list, accent_hex: str, font_name: str):
    """Parsea y renderiza tabla markdown."""
    accent_rgb = hex_to_rgb(accent_hex)
    dark_navy = "1A1A2E"

    # Filtrar líneas de separación
    rows = [line for line in table_lines if not re.match(r'^\|[-:\s|]+\|$', line.strip())]
    if not rows:
        return

    # Parsear celdas
    parsed_rows = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    num_cols = max(len(r) for r in parsed_rows)
    tbl = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    tbl.style = "Table Grid"
    tbl.autofit = True

    for i, row_data in enumerate(parsed_rows):
        row_obj = tbl.rows[i]
        is_header = (i == 0)

        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row_obj.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if is_header:
                set_cell_bg(cell, dark_navy)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(cell_text)
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                bg = "F8F9FA" if i % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                set_cell_border(cell)
                p = cell.paragraphs[0]
                apply_inline_formatting(
                    cell_text, p, font_name, 10, accent_rgb, (40, 40, 60)
                )

    # Espaciado post-tabla
    doc.add_paragraph().space_after = Pt(4)


# ─── Convertidor principal ───────────────────────────────────────────────────

def md_to_docx(md_path: str, output_path: str, title: str, subtitle: str,
               company: str, project: str, accent_hex: str, logo_path: str = None):
    """Convierte un archivo markdown a .docx."""
    accent_rgb = hex_to_rgb(accent_hex)
    dark_navy_rgb = hex_to_rgb("1A1A2E")
    font_name = "Calibri"

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Portada
    add_cover_page(doc, title, subtitle, company, accent_hex, logo_path)

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # ── Bloque de código ────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                # Inicio de bloque de código
                if in_table and table_lines:
                    parse_table(doc, table_lines, accent_hex, font_name)
                    table_lines = []
                    in_table = False

                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                # Fin de bloque de código
                in_code_block = False
                code_text = "\n".join(code_lines)

                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

                # Fondo gris claro
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F0F2F5")
                pPr.append(shd)

                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 30, 50)

                # Etiqueta de lenguaje
                if code_lang:
                    lang_p = doc.add_paragraph()
                    lang_run = lang_p.add_run(f"  {code_lang}")
                    lang_run.font.name = font_name
                    lang_run.font.size = Pt(8)
                    lang_run.font.color.rgb = RGBColor(*accent_rgb)
                    lang_p.paragraph_format.space_after = Pt(8)

                code_lines = []
                code_lang = ""
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Tabla ────────────────────────────────────────────────────────────
        if "|" in line and line.strip().startswith("|"):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            parse_table(doc, table_lines, accent_hex, font_name)
            table_lines = []
            in_table = False

        stripped = line.strip()

        # ── Línea vacía ──────────────────────────────────────────────────────
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        h4 = re.match(r'^#{4}\s+(.*)', stripped)
        h3 = re.match(r'^#{3}\s+(.*)', stripped)
        h2 = re.match(r'^#{2}\s+(.*)', stripped)
        h1 = re.match(r'^#{1}\s+(.*)', stripped)

        if h4:
            p = doc.add_paragraph()
            run = p.add_run(h4.group(1))
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(*accent_rgb)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif h3:
            p = doc.add_paragraph()
            run = p.add_run(h3.group(1))
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(*dark_navy_rgb)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif h2:
            p = doc.add_paragraph()
            run = p.add_run(h2.group(1).upper())
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(*dark_navy_rgb)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            # Línea de acento bajo H2
            accent_p = doc.add_paragraph()
            accent_run = accent_p.add_run("─" * 40)
            accent_run.font.color.rgb = RGBColor(*accent_rgb)
            accent_run.font.size = Pt(8)
            accent_p.paragraph_format.space_after = Pt(6)
        elif h1:
            p = doc.add_paragraph()
            run = p.add_run(h1.group(1).upper())
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(*dark_navy_rgb)
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(10)

        # ── Listas ────────────────────────────────────────────────────────────
        elif re.match(r'^(\s*)[-*+]\s+(.*)', line):
            m = re.match(r'^(\s*)[-*+]\s+(.*)', line)
            indent_level = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(2)
            apply_inline_formatting(
                m.group(2), p, font_name, 11, accent_rgb
            )

        elif re.match(r'^(\s*)\d+\.\s+(.*)', line):
            m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
            indent_level = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(2)
            apply_inline_formatting(
                m.group(2), p, font_name, 11, accent_rgb
            )

        # ── Blockquote ────────────────────────────────────────────────────────
        elif stripped.startswith(">"):
            text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Borde izquierdo de acento
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "16")
            left.set(qn("w:color"), accent_hex.lstrip("#"))
            pBdr.append(left)
            pPr.append(pBdr)
            apply_inline_formatting(text, p, font_name, 11, accent_rgb, (80, 80, 100))

        # ── Regla horizontal ─────────────────────────────────────────────────
        elif re.match(r'^[-*_]{3,}$', stripped):
            p = doc.add_paragraph()
            run = p.add_run("─" * 70)
            run.font.color.rgb = RGBColor(200, 200, 210)
            run.font.size = Pt(8)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)

        # ── Advertencia / nota especial ──────────────────────────────────────
        elif stripped.startswith("⚠️") or stripped.upper().startswith("ADVERTENCIA:"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "FFF8E1")
            pPr.append(shd)
            apply_inline_formatting(stripped, p, font_name, 11, accent_rgb, (120, 80, 0))

        # ── Párrafo normal ───────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            apply_inline_formatting(stripped, p, font_name, 11, accent_rgb)

        i += 1

    # Procesar tabla pendiente al final
    if in_table and table_lines:
        parse_table(doc, table_lines, accent_hex, font_name)

    doc.save(output_path)
    print(f"  ✅ {os.path.basename(output_path)}")


# ─── Mapeado de archivos md → título/subtítulo ──────────────────────────────

FILE_TITLES = {
    "README": ("Documentación General", "Descripción del proyecto"),
    "GUIA-DE-USO": ("Guía de Uso", "Manual para el usuario"),
    "DOCS-TECNICA": ("Documentación Técnica", "Arquitectura y referencia técnica"),
    "DEPLOYMENT": ("Guía de Despliegue", "Instrucciones de infraestructura y CI/CD"),
    "ACCESOS": ("Accesos del Proyecto", "DOCUMENTO CONFIDENCIAL"),
}


# ─── CLI ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convierte archivos Markdown a .docx con diseño profesional"
    )
    parser.add_argument("docs_dir", nargs="?", default="docs/entrega",
                        help="Directorio con los archivos .md (default: docs/entrega)")
    parser.add_argument("--color", default="#00E5FF",
                        help="Color de acento hex (default: #00E5FF)")
    parser.add_argument("--company", default="",
                        help="Nombre de la empresa/desarrollador")
    parser.add_argument("--logo", default=None,
                        help="Ruta al logo (PNG/JPG)")
    parser.add_argument("--project", default="Proyecto",
                        help="Nombre del proyecto")
    args = parser.parse_args()

    docs_dir = Path(args.docs_dir)
    if not docs_dir.exists():
        print(f"ERROR: Directorio '{docs_dir}' no encontrado.")
        sys.exit(1)

    md_files = {
        "README": docs_dir / "README.md",
        "GUIA-DE-USO": docs_dir / "GUIA-DE-USO.md",
        "DOCS-TECNICA": docs_dir / "DOCS-TECNICA.md",
        "DEPLOYMENT": docs_dir / "DEPLOYMENT.md",
        "ACCESOS": docs_dir / "ACCESOS.md",
    }

    print(f"\n📄 Convirtiendo documentos en: {docs_dir}")
    print(f"   Proyecto: {args.project}")
    print(f"   Color: {args.color}\n")

    converted = 0
    for key, md_path in md_files.items():
        if not md_path.exists():
            print(f"  ⚠️  {md_path.name} no encontrado, omitiendo.")
            continue

        title_info = FILE_TITLES.get(key, (args.project, ""))
        full_title = f"{args.project}\n{title_info[0]}"
        output_path = str(docs_dir / f"{key}.docx")

        try:
            md_to_docx(
                md_path=str(md_path),
                output_path=output_path,
                title=f"{args.project} — {title_info[0]}",
                subtitle=title_info[1],
                company=args.company,
                project=args.project,
                accent_hex=args.color,
                logo_path=args.logo,
            )
            converted += 1
        except Exception as e:
            print(f"  ❌ Error convirtiendo {md_path.name}: {e}")

    print(f"\n✅ {converted} documento(s) convertido(s) a Word en {docs_dir}/")


if __name__ == "__main__":
    main()
