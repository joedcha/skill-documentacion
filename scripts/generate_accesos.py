#!/usr/bin/env python3
"""
generate_accesos.py — Genera ACCESOS-DEL-PROYECTO.docx con diseño profesional
Skill: /documentacion v2.0.0
"""

import sys
import os
import re
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx no está instalado.")
    print("Instalar con: pip install python-docx")
    sys.exit(1)


# ─── Catálogo de servicios ───────────────────────────────────────────────────

SERVICES_CATALOG = {
    # Bases de datos
    "supabase": {
        "name": "Supabase",
        "category": "Base de Datos / Auth",
        "dashboard": "https://supabase.com/dashboard",
        "env_vars": ["SUPABASE_URL", "SUPABASE_ANON_KEY", "SUPABASE_SERVICE_ROLE_KEY"],
        "purpose": "Base de datos PostgreSQL, autenticación y storage",
    },
    "neon": {
        "name": "Neon",
        "category": "Base de Datos",
        "dashboard": "https://console.neon.tech",
        "env_vars": ["DATABASE_URL"],
        "purpose": "Base de datos PostgreSQL serverless",
    },
    "planetscale": {
        "name": "PlanetScale",
        "category": "Base de Datos",
        "dashboard": "https://app.planetscale.com",
        "env_vars": ["DATABASE_URL"],
        "purpose": "Base de datos MySQL serverless",
    },
    "mongodb": {
        "name": "MongoDB Atlas",
        "category": "Base de Datos",
        "dashboard": "https://cloud.mongodb.com",
        "env_vars": ["MONGODB_URI"],
        "purpose": "Base de datos NoSQL en la nube",
    },
    "convex": {
        "name": "Convex",
        "category": "Base de Datos reactiva",
        "dashboard": "https://dashboard.convex.dev",
        "env_vars": ["CONVEX_DEPLOYMENT", "NEXT_PUBLIC_CONVEX_URL"],
        "purpose": "Base de datos reactiva en tiempo real",
    },
    "firebase": {
        "name": "Firebase",
        "category": "Base de Datos / Auth / Storage",
        "dashboard": "https://console.firebase.google.com",
        "env_vars": ["FIREBASE_API_KEY", "FIREBASE_PROJECT_ID", "FIREBASE_APP_ID"],
        "purpose": "Plataforma de desarrollo de Google (DB, Auth, Storage)",
    },
    "turso": {
        "name": "Turso",
        "category": "Base de Datos",
        "dashboard": "https://app.turso.tech",
        "env_vars": ["TURSO_DATABASE_URL", "TURSO_AUTH_TOKEN"],
        "purpose": "Base de datos SQLite distribuida en el edge",
    },
    "upstash": {
        "name": "Upstash",
        "category": "Cache / Redis",
        "dashboard": "https://console.upstash.com",
        "env_vars": ["UPSTASH_REDIS_REST_URL", "UPSTASH_REDIS_REST_TOKEN"],
        "purpose": "Redis serverless para caché y rate limiting",
    },
    "pinecone": {
        "name": "Pinecone",
        "category": "Base de Datos Vectorial",
        "dashboard": "https://app.pinecone.io",
        "env_vars": ["PINECONE_API_KEY", "PINECONE_ENVIRONMENT"],
        "purpose": "Base de datos vectorial para búsqueda semántica e IA",
    },
    # Autenticación
    "clerk": {
        "name": "Clerk",
        "category": "Autenticación",
        "dashboard": "https://dashboard.clerk.com",
        "env_vars": ["CLERK_SECRET_KEY", "NEXT_PUBLIC_CLERK_PUBLISHABLE_KEY"],
        "purpose": "Sistema completo de autenticación y gestión de usuarios",
    },
    "auth0": {
        "name": "Auth0",
        "category": "Autenticación",
        "dashboard": "https://manage.auth0.com",
        "env_vars": ["AUTH0_SECRET", "AUTH0_BASE_URL", "AUTH0_ISSUER_BASE_URL", "AUTH0_CLIENT_ID", "AUTH0_CLIENT_SECRET"],
        "purpose": "Plataforma de identidad y autenticación",
    },
    "nextauth": {
        "name": "NextAuth / Auth.js",
        "category": "Autenticación",
        "dashboard": "https://authjs.dev",
        "env_vars": ["NEXTAUTH_SECRET", "NEXTAUTH_URL"],
        "purpose": "Autenticación para aplicaciones Next.js",
    },
    "better-auth": {
        "name": "Better Auth",
        "category": "Autenticación",
        "dashboard": "https://better-auth.com",
        "env_vars": ["BETTER_AUTH_SECRET", "BETTER_AUTH_URL"],
        "purpose": "Framework de autenticación flexible",
    },
    # Hosting
    "vercel": {
        "name": "Vercel",
        "category": "Hosting / Deploy",
        "dashboard": "https://vercel.com/dashboard",
        "env_vars": [],
        "purpose": "Plataforma de hosting y despliegue (configurar desde dashboard)",
    },
    "netlify": {
        "name": "Netlify",
        "category": "Hosting / Deploy",
        "dashboard": "https://app.netlify.com",
        "env_vars": [],
        "purpose": "Plataforma de hosting y despliegue (configurar desde dashboard)",
    },
    "railway": {
        "name": "Railway",
        "category": "Hosting / Deploy",
        "dashboard": "https://railway.app/dashboard",
        "env_vars": [],
        "purpose": "Plataforma de hosting con base de datos integrada",
    },
    "render": {
        "name": "Render",
        "category": "Hosting / Deploy",
        "dashboard": "https://dashboard.render.com",
        "env_vars": [],
        "purpose": "Plataforma de hosting para apps y bases de datos",
    },
    "fly": {
        "name": "Fly.io",
        "category": "Hosting / Deploy",
        "dashboard": "https://fly.io/dashboard",
        "env_vars": ["FLY_API_TOKEN"],
        "purpose": "Plataforma de hosting distribuido globalmente",
    },
    "cloudflare": {
        "name": "Cloudflare",
        "category": "DNS / CDN / Pages",
        "dashboard": "https://dash.cloudflare.com",
        "env_vars": ["CLOUDFLARE_API_TOKEN", "CLOUDFLARE_ACCOUNT_ID"],
        "purpose": "DNS, CDN, protección DDoS y hosting de workers/pages",
    },
    "aws": {
        "name": "Amazon Web Services",
        "category": "Cloud / Infraestructura",
        "dashboard": "https://console.aws.amazon.com",
        "env_vars": ["AWS_ACCESS_KEY_ID", "AWS_SECRET_ACCESS_KEY", "AWS_REGION"],
        "purpose": "Infraestructura cloud (S3, Lambda, RDS, etc.)",
    },
    "gcp": {
        "name": "Google Cloud Platform",
        "category": "Cloud / Infraestructura",
        "dashboard": "https://console.cloud.google.com",
        "env_vars": ["GOOGLE_APPLICATION_CREDENTIALS", "GCP_PROJECT_ID"],
        "purpose": "Infraestructura cloud de Google",
    },
    # IA
    "openai": {
        "name": "OpenAI",
        "category": "IA / LLM",
        "dashboard": "https://platform.openai.com",
        "env_vars": ["OPENAI_API_KEY"],
        "purpose": "API de modelos GPT y embeddings",
    },
    "anthropic": {
        "name": "Anthropic",
        "category": "IA / LLM",
        "dashboard": "https://console.anthropic.com",
        "env_vars": ["ANTHROPIC_API_KEY"],
        "purpose": "API de modelos Claude",
    },
    "replicate": {
        "name": "Replicate",
        "category": "IA / Modelos",
        "dashboard": "https://replicate.com/account",
        "env_vars": ["REPLICATE_API_TOKEN"],
        "purpose": "API para ejecutar modelos de IA open source",
    },
    "llamaindex": {
        "name": "LlamaIndex / LlamaParse",
        "category": "IA / Document Parsing",
        "dashboard": "https://cloud.llamaindex.ai",
        "env_vars": ["LLAMA_CLOUD_API_KEY"],
        "purpose": "Parseo inteligente de documentos con IA",
    },
    # Pagos
    "stripe": {
        "name": "Stripe",
        "category": "Pagos",
        "dashboard": "https://dashboard.stripe.com",
        "env_vars": ["STRIPE_SECRET_KEY", "STRIPE_WEBHOOK_SECRET", "NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY"],
        "purpose": "Procesamiento de pagos, suscripciones y facturas",
    },
    "mercadopago": {
        "name": "Mercado Pago",
        "category": "Pagos",
        "dashboard": "https://www.mercadopago.com.ar/developers",
        "env_vars": ["MP_ACCESS_TOKEN", "MP_PUBLIC_KEY"],
        "purpose": "Procesamiento de pagos para Latinoamérica",
    },
    "paypal": {
        "name": "PayPal",
        "category": "Pagos",
        "dashboard": "https://developer.paypal.com/dashboard",
        "env_vars": ["PAYPAL_CLIENT_ID", "PAYPAL_CLIENT_SECRET"],
        "purpose": "Procesamiento de pagos globales",
    },
    # Email
    "resend": {
        "name": "Resend",
        "category": "Email",
        "dashboard": "https://resend.com/overview",
        "env_vars": ["RESEND_API_KEY"],
        "purpose": "Envío de emails transaccionales",
    },
    "sendgrid": {
        "name": "SendGrid",
        "category": "Email",
        "dashboard": "https://app.sendgrid.com",
        "env_vars": ["SENDGRID_API_KEY"],
        "purpose": "Envío de emails transaccionales y marketing",
    },
    "mailgun": {
        "name": "Mailgun",
        "category": "Email",
        "dashboard": "https://app.mailgun.com",
        "env_vars": ["MAILGUN_API_KEY", "MAILGUN_DOMAIN"],
        "purpose": "Envío de emails transaccionales con alta entregabilidad",
    },
    "postmark": {
        "name": "Postmark",
        "category": "Email",
        "dashboard": "https://account.postmarkapp.com",
        "env_vars": ["POSTMARK_SERVER_TOKEN"],
        "purpose": "Envío de emails transaccionales de alta velocidad",
    },
    # Comunicación
    "twilio": {
        "name": "Twilio",
        "category": "SMS / Comunicación",
        "dashboard": "https://console.twilio.com",
        "env_vars": ["TWILIO_ACCOUNT_SID", "TWILIO_AUTH_TOKEN", "TWILIO_PHONE_NUMBER"],
        "purpose": "Envío de SMS, llamadas y comunicaciones programáticas",
    },
    # Storage / Media
    "uploadthing": {
        "name": "UploadThing",
        "category": "Storage / Archivos",
        "dashboard": "https://uploadthing.com/dashboard",
        "env_vars": ["UPLOADTHING_SECRET", "UPLOADTHING_APP_ID"],
        "purpose": "Carga y gestión de archivos para aplicaciones Next.js",
    },
    "cloudinary": {
        "name": "Cloudinary",
        "category": "Media / Storage",
        "dashboard": "https://cloudinary.com/console",
        "env_vars": ["CLOUDINARY_URL", "CLOUDINARY_CLOUD_NAME", "CLOUDINARY_API_KEY", "CLOUDINARY_API_SECRET"],
        "purpose": "Gestión, optimización y entrega de imágenes y videos",
    },
    "s3": {
        "name": "Amazon S3",
        "category": "Storage",
        "dashboard": "https://s3.console.aws.amazon.com",
        "env_vars": ["AWS_S3_BUCKET", "AWS_ACCESS_KEY_ID", "AWS_SECRET_ACCESS_KEY", "AWS_REGION"],
        "purpose": "Almacenamiento de archivos y objetos en la nube",
    },
    # Monitoreo
    "sentry": {
        "name": "Sentry",
        "category": "Monitoreo / Errores",
        "dashboard": "https://sentry.io/organizations/",
        "env_vars": ["SENTRY_DSN", "NEXT_PUBLIC_SENTRY_DSN", "SENTRY_AUTH_TOKEN"],
        "purpose": "Monitoreo de errores y performance en producción",
    },
    "posthog": {
        "name": "PostHog",
        "category": "Analytics / Product",
        "dashboard": "https://app.posthog.com",
        "env_vars": ["POSTHOG_API_KEY", "NEXT_PUBLIC_POSTHOG_KEY", "NEXT_PUBLIC_POSTHOG_HOST"],
        "purpose": "Analytics de producto, feature flags y grabaciones de sesión",
    },
    "datadog": {
        "name": "Datadog",
        "category": "Monitoreo / APM",
        "dashboard": "https://app.datadoghq.com",
        "env_vars": ["DD_API_KEY", "DD_APP_KEY"],
        "purpose": "Monitoreo de infraestructura, logs y APM",
    },
    "betterstack": {
        "name": "Better Stack / Logtail",
        "category": "Logs / Uptime",
        "dashboard": "https://betterstack.com/dashboard",
        "env_vars": ["LOGTAIL_SOURCE_TOKEN", "BETTER_STACK_SOURCE_TOKEN"],
        "purpose": "Gestión de logs y monitoreo de uptime",
    },
    # Repositorio
    "github": {
        "name": "GitHub",
        "category": "Repositorio / CI/CD",
        "dashboard": "https://github.com",
        "env_vars": ["GITHUB_TOKEN", "GH_TOKEN"],
        "purpose": "Repositorio de código fuente y CI/CD con GitHub Actions",
    },
    "gitlab": {
        "name": "GitLab",
        "category": "Repositorio / CI/CD",
        "dashboard": "https://gitlab.com",
        "env_vars": ["GITLAB_TOKEN"],
        "purpose": "Repositorio de código fuente y CI/CD",
    },
}

# Aliases para normalizar nombres de servicio
SERVICE_ALIASES = {
    "next-auth": "nextauth",
    "better_auth": "better-auth",
    "mongo": "mongodb",
    "mongo-atlas": "mongodb",
    "fly.io": "fly",
    "llamaparse": "llamaindex",
    "llama": "llamaindex",
    "s3": "s3",
    "logtail": "betterstack",
    "better-stack": "betterstack",
}


# ─── Utilidades de formato ───────────────────────────────────────────────────

def hex_to_rgb(hex_color: str) -> tuple:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return (0, 229, 255)
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def add_heading(doc, text: str, level: int, accent_rgb: tuple, dark_rgb: tuple,
                font_name: str = "Calibri"):
    p = doc.add_paragraph()
    run = p.add_run(text.upper() if level <= 2 else text)
    run.bold = True
    run.font.name = font_name

    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(*dark_rgb)
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*dark_rgb)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        # Línea de acento
        ap = doc.add_paragraph()
        ar = ap.add_run("─" * 40)
        ar.font.color.rgb = RGBColor(*accent_rgb)
        ar.font.size = Pt(8)
        ap.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(*accent_rgb)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    return p


def add_info_table(doc, rows: list, accent_hex: str, font_name: str = "Calibri"):
    """Crea tabla de dos columnas: Campo | Valor."""
    dark_navy = "1A1A2E"
    accent_rgb = hex_to_rgb(accent_hex)

    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.style = "Table Grid"

    # Encabezado
    header_row = tbl.rows[0]
    for j, header in enumerate(["Campo", "Valor"]):
        cell = header_row.cells[j]
        set_cell_bg(cell, dark_navy)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Filas de datos
    for i, (campo, valor) in enumerate(rows):
        row_obj = tbl.rows[i + 1]
        bg = "F8F9FA" if i % 2 == 0 else "FFFFFF"

        # Campo
        c0 = row_obj.cells[0]
        set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(campo)
        r0.bold = True
        r0.font.name = font_name
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(30, 30, 50)

        # Valor
        c1 = row_obj.cells[1]
        set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(valor)
        r1.font.name = font_name
        r1.font.size = Pt(10)
        if valor.startswith("["):
            r1.font.color.rgb = RGBColor(150, 100, 0)
        else:
            r1.font.color.rgb = RGBColor(40, 40, 60)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_cover_page(doc, project: str, client: str, company: str,
                   accent_hex: str, logo_path: str = None):
    """Portada del documento de accesos."""
    accent_rgb = hex_to_rgb(accent_hex)
    dark_rgb = hex_to_rgb("1A1A2E")
    font_name = "Calibri"

    if logo_path and os.path.exists(logo_path):
        try:
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lp.add_run().add_picture(logo_path, width=Inches(1.8))
            lp.space_after = Pt(30)
        except Exception:
            pass

    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Badge confidencial
    badge_p = doc.add_paragraph()
    badge_run = badge_p.add_run("⚠  DOCUMENTO CONFIDENCIAL")
    badge_run.bold = True
    badge_run.font.name = font_name
    badge_run.font.size = Pt(10)
    badge_run.font.color.rgb = RGBColor(180, 50, 0)
    badge_p.space_after = Pt(20)

    # Línea de acento
    ap = doc.add_paragraph()
    ar = ap.add_run("█" * 6)
    ar.font.color.rgb = RGBColor(*accent_rgb)
    ar.font.size = Pt(14)
    ap.space_after = Pt(12)

    # Título
    tp = doc.add_paragraph()
    tr = tp.add_run(f"{project.upper()}")
    tr.bold = True
    tr.font.name = font_name
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(*dark_rgb)
    tp.space_after = Pt(6)

    # Subtítulo
    stp = doc.add_paragraph()
    str_ = stp.add_run("Accesos y Credenciales del Proyecto")
    str_.font.name = font_name
    str_.font.size = Pt(14)
    str_.font.color.rgb = RGBColor(80, 80, 100)
    stp.space_after = Pt(40)

    # Línea divisoria
    dp = doc.add_paragraph()
    dr = dp.add_run("─" * 52)
    dr.font.color.rgb = RGBColor(*accent_rgb)
    dr.font.size = Pt(10)
    dp.space_after = Pt(16)

    # Info
    MONTHS_ES = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    now = datetime.now()
    date_str = f"{now.day} de {MONTHS_ES[now.month]} de {now.year}"

    for label, value in [("Cliente", client), ("Preparado por", company), ("Fecha", date_str)]:
        ip = doc.add_paragraph()
        ir_label = ip.add_run(f"{label}: ")
        ir_label.bold = True
        ir_label.font.name = font_name
        ir_label.font.size = Pt(11)
        ir_label.font.color.rgb = RGBColor(80, 80, 100)
        ir_value = ip.add_run(value)
        ir_value.font.name = font_name
        ir_value.font.size = Pt(11)
        ir_value.font.color.rgb = RGBColor(30, 30, 50)
        ip.space_after = Pt(4)

    doc.add_page_break()


# ─── Generador principal ─────────────────────────────────────────────────────

def generate_accesos(project: str, client: str, company: str, accent_hex: str,
                     services_input: str, output_path: str, logo_path: str = None):
    """Genera el documento completo de accesos."""
    accent_rgb = hex_to_rgb(accent_hex)
    dark_rgb = hex_to_rgb("1A1A2E")
    font_name = "Calibri"

    # Normalizar lista de servicios
    raw_services = [s.strip().lower() for s in services_input.split(",") if s.strip()]
    resolved_services = []
    for svc in raw_services:
        key = SERVICE_ALIASES.get(svc, svc)
        if key in SERVICES_CATALOG:
            resolved_services.append(key)
        else:
            # Búsqueda parcial
            for catalog_key in SERVICES_CATALOG:
                if svc in catalog_key or catalog_key in svc:
                    resolved_services.append(catalog_key)
                    break

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Portada
    add_cover_page(doc, project, client, company, accent_hex, logo_path)

    # ── Aviso de seguridad ───────────────────────────────────────────────────
    add_heading(doc, "Aviso de Seguridad", 2, accent_rgb, dark_rgb, font_name)

    security_notes = [
        "NO compartir este documento por canales no seguros (email sin cifrar, WhatsApp, Slack).",
        "Guardar en un gestor de contraseñas: 1Password, Bitwarden, Dashlane o equivalente.",
        "Cambiar todas las contraseñas dentro de los primeros 7 días de recibidas.",
        "Revocar los accesos del desarrollador una vez confirmada la recepción.",
        "Nunca commitear credenciales reales al repositorio de código.",
    ]
    for note in security_notes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"• {note}")
        run.font.name = font_name
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(120, 40, 0)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Información general ──────────────────────────────────────────────────
    add_heading(doc, "Información General", 2, accent_rgb, dark_rgb, font_name)
    add_info_table(doc, [
        ("Proyecto", project),
        ("Cliente", client),
        ("Desarrollado por", company),
        ("URL de producción", "[A_COMPLETAR]"),
        ("Repositorio de código", "[A_COMPLETAR]"),
        ("Branch principal", "main"),
        ("Fecha de entrega", datetime.now().strftime("%d/%m/%Y")),
    ], accent_hex, font_name)

    # ── Acceso a la aplicación ───────────────────────────────────────────────
    add_heading(doc, "Acceso a la Aplicación Web", 2, accent_rgb, dark_rgb, font_name)
    add_info_table(doc, [
        ("URL de producción", "[A_COMPLETAR]"),
        ("Email administrador", "[A_COMPLETAR]"),
        ("Contraseña inicial", "[CAMBIAR INMEDIATAMENTE]"),
        ("Panel de administración", "[A_COMPLETAR si existe]"),
    ], accent_hex, font_name)

    # ── Secciones por categoría ──────────────────────────────────────────────
    categories_order = [
        "Base de Datos", "Base de Datos / Auth", "Base de Datos reactiva",
        "Base de Datos Vectorial", "Cache / Redis",
        "Autenticación",
        "Hosting / Deploy", "DNS / CDN / Pages", "Cloud / Infraestructura",
        "Storage / Archivos", "Storage", "Media / Storage",
        "IA / LLM", "IA / Modelos", "IA / Document Parsing",
        "Pagos",
        "Email",
        "SMS / Comunicación",
        "Monitoreo / Errores", "Analytics / Product", "Monitoreo / APM", "Logs / Uptime",
        "Repositorio / CI/CD",
    ]

    # Agrupar servicios detectados por categoría
    by_category = {}
    for svc_key in resolved_services:
        svc = SERVICES_CATALOG[svc_key]
        cat = svc["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append((svc_key, svc))

    # Ordenar por categoría predefinida
    sorted_categories = sorted(
        by_category.keys(),
        key=lambda c: categories_order.index(c) if c in categories_order else 999
    )

    for cat in sorted_categories:
        svcs = by_category[cat]
        add_heading(doc, cat, 2, accent_rgb, dark_rgb, font_name)

        for svc_key, svc in svcs:
            add_heading(doc, svc["name"], 3, accent_rgb, dark_rgb, font_name)

            # Info básica
            info_rows = [
                ("Dashboard", svc["dashboard"]),
                ("Propósito", svc["purpose"]),
                ("Email / Usuario", "[A_COMPLETAR]"),
                ("Contraseña / API Key", "[A_COMPLETAR]"),
            ]

            if svc["env_vars"]:
                env_str = ", ".join(f"`{v}`" for v in svc["env_vars"])
                info_rows.append(("Variables de entorno", env_str))

            info_rows.append(("Notas adicionales", "[A_COMPLETAR si aplica]"))

            add_info_table(doc, info_rows, accent_hex, font_name)

    # ── Variables de entorno ─────────────────────────────────────────────────
    add_heading(doc, "Variables de Entorno (.env)", 2, accent_rgb, dark_rgb, font_name)

    p = doc.add_paragraph()
    run = p.add_run(
        "El archivo .env con los valores reales de producción está guardado en:"
    )
    run.font.name = font_name
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(60, 60, 80)
    p.space_after = Pt(8)

    add_info_table(doc, [
        ("Ubicación del .env de producción", "[A_COMPLETAR: ej. 1Password > Proyecto X > .env prod]"),
        ("Referencia en repositorio", ".env.example (sin valores reales)"),
        ("Variables totales documentadas", f"{sum(len(SERVICES_CATALOG[k]['env_vars']) for k in resolved_services if k in SERVICES_CATALOG)} variables"),
    ], accent_hex, font_name)

    # ── Repositorio de código ────────────────────────────────────────────────
    add_heading(doc, "Repositorio de Código", 2, accent_rgb, dark_rgb, font_name)
    add_info_table(doc, [
        ("Plataforma", "[GitHub / GitLab / Bitbucket]"),
        ("URL del repositorio", "[A_COMPLETAR]"),
        ("Branch principal", "main"),
        ("Acceso otorgado a", "[EMAIL DEL CLIENTE]"),
        ("Tipo de acceso", "[Owner / Admin / Write / Read]"),
    ], accent_hex, font_name)

    # ── Confirmación de recepción ────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Confirmación de Recepción", 2, accent_rgb, dark_rgb, font_name)

    p = doc.add_paragraph()
    run = p.add_run(
        "El cliente declara haber recibido y verificado todos los accesos listados "
        "en este documento, y asume la responsabilidad sobre su custodia segura."
    )
    run.font.name = font_name
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(60, 60, 80)
    p.space_after = Pt(24)

    # Firma
    for _ in range(3):
        doc.add_paragraph()

    firma_rows = [
        ("Nombre completo", "_" * 40),
        ("Firma", "_" * 40),
        ("Fecha", "_" * 40),
        ("Cargo / Rol", "_" * 40),
    ]
    for campo, linea in firma_rows:
        fp = doc.add_paragraph()
        fp.paragraph_format.space_after = Pt(16)
        fr1 = fp.add_run(f"{campo}: ")
        fr1.bold = True
        fr1.font.name = font_name
        fr1.font.size = Pt(11)
        fr1.font.color.rgb = RGBColor(60, 60, 80)
        fr2 = fp.add_run(linea)
        fr2.font.name = font_name
        fr2.font.size = Pt(11)
        fr2.font.color.rgb = RGBColor(180, 180, 200)

    doc.save(output_path)
    print(f"  ✅ {os.path.basename(output_path)}")
    print(f"     Servicios documentados: {len(resolved_services)}")
    if resolved_services:
        for svc_key in resolved_services:
            print(f"       • {SERVICES_CATALOG[svc_key]['name']}")


# ─── CLI ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Genera ACCESOS-DEL-PROYECTO.docx con diseño profesional"
    )
    parser.add_argument("--project", default="Proyecto", help="Nombre del proyecto")
    parser.add_argument("--client", default="Cliente", help="Nombre del cliente")
    parser.add_argument("--company", default="", help="Nombre de la empresa/desarrollador")
    parser.add_argument("--color", default="#00E5FF", help="Color de acento hex")
    parser.add_argument("--logo", default=None, help="Ruta al logo (PNG/JPG)")
    parser.add_argument("--services", default="",
                        help="Servicios detectados, separados por comas (ej: supabase,stripe,clerk)")
    parser.add_argument("--output", default="docs/entrega/ACCESOS-DEL-PROYECTO.docx",
                        help="Ruta de salida del .docx")
    args = parser.parse_args()

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"\n🔐 Generando documento de accesos...")
    print(f"   Proyecto: {args.project}")
    print(f"   Cliente: {args.client}")
    print(f"   Servicios: {args.services or '(ninguno especificado)'}\n")

    generate_accesos(
        project=args.project,
        client=args.client,
        company=args.company,
        accent_hex=args.color,
        services_input=args.services,
        output_path=str(output_path),
        logo_path=args.logo,
    )

    print(f"\n✅ Documento guardado en: {output_path}")
    print("   ⚠️  Completar los campos [A_COMPLETAR] antes de entregar al cliente.")


if __name__ == "__main__":
    main()
